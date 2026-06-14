from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = r"C:\Users\User\AndroidStudioProjects\Lumiere\Lumiere_Portfolio_Tools_AI.docx"

GOLD = RGBColor(201, 162, 109)
DARK = RGBColor(44, 26, 10)
MUTED = RGBColor(122, 96, 69)
LIGHT = RGBColor(234, 216, 192)
WHITE = RGBColor(255, 255, 255)
GREEN = RGBColor(64, 160, 112)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=110, start=110, bottom=110, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name, size=None, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def style_paragraph(paragraph, space_before=0, space_after=0, line=1.15, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(paragraph, text, font="Calibri", size=11, bold=False, color=DARK, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, font, size=size, bold=bold, color=color, italic=italic)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=2, space_after=8, line=1.0)
    add_text(p, text, font="Cambria", size=18 if level == 1 else 14, bold=True, color=DARK)
    if level == 1:
        p_format = p.paragraph_format
        p_format.keep_with_next = True
    return p


def add_body(doc, text, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    style_paragraph(p, space_after=5, line=1.12)
    add_text(p, text, font="Calibri", size=10.5, color=DARK)
    return p


def add_note_box(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F9F3EA")
    set_cell_margins(cell, top=140, start=150, bottom=140, end=150)
    p = cell.paragraphs[0]
    style_paragraph(p, space_after=4)
    add_text(p, title, font="Cambria", size=11.5, bold=True, color=DARK)
    for line in lines:
        q = cell.add_paragraph()
        style_paragraph(q, space_after=3, line=1.08)
        add_text(q, line, font="Calibri", size=10, color=MUTED)
    doc.add_paragraph()


def add_table(doc, headers, rows, column_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        set_cell_shading(cell, "C9A26D")
        set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, header, font="Calibri", size=10, bold=True, color=WHITE)
        if column_widths:
            cell.width = Inches(column_widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.text = ""
            set_cell_margins(cell, top=100, start=110, bottom=100, end=110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, line=1.05)
            add_text(p, str(value), font="Calibri", size=9.5, color=DARK)
            if column_widths:
                cell.width = Inches(column_widths[i])
    doc.add_paragraph()
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        p,
        "Lumiere Portfolio | Mobile Application Development",
        font="Calibri",
        size=9,
        color=MUTED,
    )


def add_page_break(doc):
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_footer(section)


def make_cover(doc):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=30, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "Mobile Application Development", font="Cambria", size=20, bold=True, color=GOLD)

    p = doc.add_paragraph()
    style_paragraph(p, space_before=8, space_after=8, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        p,
        "Lumiere: A Salon Booking And Gifting Experience App",
        font="Cambria",
        size=24,
        bold=True,
        color=DARK,
    )

    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        p,
        "Portfolio of Plug-ins, Tools, Add-ons, AI Services, Backend Components, and Deployment Stack",
        font="Calibri",
        size=12,
        color=MUTED,
    )

    accent = doc.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = accent.cell(0, 0)
    set_cell_shading(cell, "F8F1E7")
    set_cell_margins(cell, top=220, start=180, bottom=220, end=180)
    p = cell.paragraphs[0]
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        p,
        "Assumption applied in this portfolio: the application is fully implemented, backed by Firebase, quality-tested, and deployed for production use.",
        font="Calibri",
        size=11.5,
        bold=True,
        color=DARK,
    )

    for line in [
        "Prepared By",
        "Hammad Ahmad  |  BSEF24M502",
        "Muhammad Ans  |  BSEF24M511",
        "Eden Nadeem Bhatti  |  BSEF24M549",
    ]:
        p = doc.add_paragraph()
        style_paragraph(p, space_before=10 if line == "Prepared By" else 4, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(
            p,
            line,
            font="Calibri",
            size=13 if line == "Prepared By" else 12,
            bold=line == "Prepared By",
            color=DARK if line != "Prepared By" else GOLD,
        )

    p = doc.add_paragraph()
    style_paragraph(p, space_before=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        p,
        "Submitted as a project portfolio for the Lumiere mobile application initiative.",
        font="Calibri",
        size=11,
        color=MUTED,
        italic=True,
    )


def add_project_overview(doc):
    add_heading(doc, "1. Project Overview and Portfolio Scope")
    add_body(doc, "Lumiere is envisioned as a premium salon booking and gifting application that allows users to discover beauty service providers, reserve appointments, send gift experiences, and manage their profiles from a single Android application.")
    add_body(doc, "For this portfolio, the project is described as a complete production-ready solution rather than only a prototype. This means the document includes the frontend mobile application, the Firebase-powered backend, deployment workflow, testing stack, analytics, security controls, monitoring, and AI-assisted development support.")
    add_body(doc, "The purpose of this report is to identify every major plug-in, tool, add-on, platform, and AI service that contributed to planning, building, testing, deploying, and maintaining the Lumiere app.")

    add_note_box(
        doc,
        "Portfolio Objective",
        [
            "Document the technical ecosystem behind the app.",
            "Show how design, development, backend, QA, and deployment were connected.",
            "Present the tools in a structured academic portfolio format.",
        ],
    )

    add_table(
        doc,
        ["Area", "What It Covered in Lumiere", "Outcome"],
        [
            ["Frontend", "Android UI, navigation, booking flows, gifting screens, user profile", "User-facing mobile experience"],
            ["Backend", "Authentication, data storage, notifications, cloud logic, media storage", "Operational Firebase backend"],
            ["Deployment", "Testing releases, production rollout, monitoring, updates", "Live and maintainable application"],
            ["AI Support", "Design assistance, coding support, issue resolution, content drafting", "Faster and more consistent delivery"],
        ],
        column_widths=[1.3, 3.8, 1.4],
    )

    add_body(doc, "Because the project team worked in a modern tool-assisted workflow, the technology mix included both conventional development utilities and AI-enabled assistants. The remaining pages break these components down in detail.")


def add_frontend_page(doc):
    add_heading(doc, "2. Frontend Mobile Application Stack")
    add_body(doc, "The Lumiere client application was developed as a native Android app. Kotlin was used as the primary programming language, while XML layout files were used to define screens such as splash, onboarding, login, home, bookings, gifting, maps, and profile.")
    add_body(doc, "Material Design components supported polished user interaction, while Glide handled remote image loading for salon cards and profile visuals. The frontend layer was responsible for navigation, form validation, state presentation, and an overall premium visual experience.")

    add_table(
        doc,
        ["Component", "Type", "Use in Lumiere"],
        [
            ["Android Studio", "IDE", "Main development environment for coding, debugging, emulation, and project builds"],
            ["Kotlin", "Language", "Core logic for activities, fragments, interaction handling, and data flow"],
            ["XML Layouts", "UI Definition", "Screen structure, cards, forms, buttons, and navigation containers"],
            ["Material Components", "UI Library", "Modern buttons, text fields, cards, switches, and navigation widgets"],
            ["Glide", "Image Library", "Loaded remote salon and avatar images efficiently"],
            ["View Binding", "Android Feature", "Safe binding between XML views and Kotlin code"],
        ],
        column_widths=[1.55, 1.2, 3.5],
    )

    add_note_box(
        doc,
        "Why This Stack Was Suitable",
        [
            "It kept the application native to Android and easy to present in Android Studio.",
            "It matched the Figma-inspired design with accurate XML control over layout and spacing.",
            "It supported dummy screens first and could scale cleanly into a production-ready app.",
        ],
    )

    add_body(doc, "From a portfolio perspective, the frontend stack was the visible face of the project. It translated design ideas into a usable Android experience and acted as the presentation layer over the assumed Firebase backend.")


def add_firebase_backend_page(doc):
    add_heading(doc, "3. Backend Architecture Using Firebase")
    add_body(doc, "For the complete deployed version of Lumiere, Firebase is assumed as the backend platform. Firebase was appropriate because it provides managed authentication, document storage, cloud functions, asset hosting support, notifications, analytics, and monitoring services within one ecosystem.")
    add_body(doc, "The backend was designed to support account creation, service listings, salon profiles, appointment bookings, gifting records, media storage, notification delivery, and real-time updates without requiring self-managed server infrastructure.")

    add_table(
        doc,
        ["Firebase Service", "Role in the System", "Lumiere Use Case"],
        [
            ["Firebase Authentication", "Identity and sign-in", "Email/password login and secure user sessions"],
            ["Cloud Firestore", "NoSQL database", "Stored salons, services, bookings, gifts, profiles, and reviews"],
            ["Cloud Functions", "Server-side logic", "Booking confirmation rules, gift issuance, and notification triggers"],
            ["Firebase Storage", "File storage", "Salon images, promotional banners, and user-uploaded media"],
            ["Firebase Cloud Messaging", "Push notification service", "Booking reminders, gift alerts, and promotional offers"],
            ["Firebase App Check", "Backend protection", "Reduced abuse from unauthorized app or script traffic"],
        ],
        column_widths=[1.7, 1.7, 2.85],
    )

    add_body(doc, "Using Firebase also simplified integration between the Android app and cloud services. Rather than building every endpoint manually, the team could rely on SDK integration, security rules, and event-driven backend functions to speed up delivery.")


def add_data_and_security_page(doc):
    add_heading(doc, "4. Data, Security, and User Management Components")
    add_body(doc, "Once Firebase is assumed as the backend, data design and security become central components of the portfolio. Lumiere would require carefully structured collections for users, salons, services, bookings, gifts, transactions, and notifications.")
    add_body(doc, "Security rules were assumed to ensure that customers could only access their own personal data while salon administrators could manage only the records assigned to them. Authentication tokens, cloud-side validation, and rules-based access would be used together.")

    add_table(
        doc,
        ["Layer", "Component Used", "Purpose"],
        [
            ["User Accounts", "Firebase Authentication", "Created and validated customer identities"],
            ["Database Rules", "Firestore Security Rules", "Controlled read/write access for sensitive app records"],
            ["Media Protection", "Firebase Storage Rules", "Secured uploaded and hosted images"],
            ["Server Validation", "Cloud Functions", "Validated booking workflows and protected key business logic"],
            ["Trust Layer", "Firebase App Check", "Prevented misuse from unofficial clients"],
            ["Recovery", "Firestore backups / exports", "Supported data continuity and operational resilience"],
        ],
        column_widths=[1.3, 2.1, 2.85],
    )

    add_note_box(
        doc,
        "Examples of Protected Data in Lumiere",
        [
            "Customer booking history and profile details",
            "Gift records and recipient information",
            "Salon service catalog and schedule management data",
            "Notification tokens and engagement events",
        ],
    )

    add_body(doc, "This layer is important in the portfolio because it shows that the project was not treated as only a UI exercise. Even under an assumed production scenario, the app architecture includes proper identity, access control, and cloud data management.")


def add_deployment_page(doc):
    add_heading(doc, "5. Deployment, Release, and Operations Stack")
    add_body(doc, "After development and backend integration, Lumiere was assumed to be deployed through a managed release pipeline. This included test distribution, production publishing, crash monitoring, analytics review, and post-release maintenance.")
    add_body(doc, "The Android app could be distributed to internal testers before production launch, while Firebase-based services would continue operating in the cloud without separate server hosting setup.")

    add_table(
        doc,
        ["Tool or Service", "Category", "Production Use"],
        [
            ["Firebase App Distribution", "Release Testing", "Shared pre-release APK or AAB builds with testers"],
            ["Google Play Console", "Store Deployment", "Published the production version of Lumiere"],
            ["Crashlytics", "Monitoring", "Tracked crashes and stability issues after release"],
            ["Firebase Analytics", "Product Insights", "Measured user behavior, booking engagement, and gift interactions"],
            ["Performance Monitoring", "Optimization", "Tracked launch speed, network delays, and screen responsiveness"],
            ["GitHub / Git", "Version Control", "Managed code history, collaboration, and release checkpoints"],
        ],
        column_widths=[1.9, 1.5, 2.85],
    )

    add_body(doc, "In a full deployment workflow, the release pipeline closes the loop between development and operations. The portfolio therefore includes not just how the app was built, but also how it was tested, observed, and maintained after going live.")


def add_design_tools_page(doc):
    add_heading(doc, "6. Design, Prototyping, and Visual Planning Tools")
    add_body(doc, "The design direction of Lumiere came from Figma-based interface generation and refinement. Figma AI supported the initial visual concept, while Figma itself served as the collaborative design platform for layouts, color palette decisions, user flow visualization, and component planning.")
    add_body(doc, "Design tools were especially important for this project because Lumiere depends heavily on premium visual presentation. The salon and gifting theme required elegant typography, warm neutral colors, and card-driven screen composition.")

    add_table(
        doc,
        ["Tool / Add-on", "Type", "Contribution to Lumiere"],
        [
            ["Figma", "Design Platform", "Created the screen structure, component hierarchy, and UI planning workspace"],
            ["Figma AI", "AI Design Assistant", "Generated early design directions and layout ideas"],
            ["Icon / Asset Export", "Design Utility", "Prepared visual elements for developer handoff"],
            ["Color and Typography Tokens", "Design System Support", "Maintained consistent styling across screens"],
        ],
        column_widths=[1.85, 1.45, 3.0],
    )

    add_note_box(
        doc,
        "Design Outputs That Informed Development",
        [
            "Splash and onboarding composition",
            "Login and profile form structure",
            "Home recommendations and salon card layout",
            "Booking, maps, and gifting feature presentation",
        ],
    )

    add_body(doc, "This page is relevant to the portfolio theme because it documents the tools and add-ons used before coding started. It shows the transition from visual concept to implementable Android screens.")


def add_ai_tools_page(doc):
    add_heading(doc, "7. AI Services, Plug-ins, and Development Assistants")
    add_body(doc, "AI support was used as an accelerator across design, coding, debugging, and content preparation. These services did not replace engineering work; instead, they reduced repetitive effort, generated alternatives quickly, and helped convert ideas into structured implementation.")
    add_body(doc, "For Lumiere, AI tools were used for UI ideation, Android code scaffolding, issue troubleshooting, resource cleanup, documentation support, and drafting technical explanations.")

    add_table(
        doc,
        ["AI Tool / Plug-in", "Where It Was Used", "Value Delivered"],
        [
            ["Gemini in Android Studio", "Coding environment", "Assisted with Android errors, Gradle guidance, and code understanding"],
            ["ChatGPT / Codex", "Development support", "Helped build screens, organize app structure, and refine implementation logic"],
            ["Figma AI", "Design phase", "Generated design concepts and screen inspiration"],
            ["AI-assisted writing support", "Documentation", "Helped draft portfolio-ready technical explanations and structure"],
        ],
        column_widths=[1.85, 1.75, 2.7],
    )

    add_body(doc, "These AI tools are central to the requested portfolio because the report specifically tracks plug-ins, tools, add-ons, and AI services. Their inclusion demonstrates how modern mobile app development increasingly combines engineering work with intelligent assistants.")


def add_qa_page(doc):
    add_heading(doc, "8. Testing, Quality Assurance, and Debugging Tools")
    add_body(doc, "A deployed application requires much more than coding and backend setup. Lumiere would have needed systematic validation of layout behavior, login flow, booking transactions, gifting logic, data synchronization, and stability across devices.")
    add_body(doc, "Testing and debugging tools supported both development-time confidence and production-time reliability.")

    add_table(
        doc,
        ["Tool", "Category", "Testing or QA Role"],
        [
            ["Android Emulator", "Execution Environment", "Validated navigation, screen layout, and functional interactions"],
            ["Physical Android Device Testing", "Real Device QA", "Confirmed usability, responsiveness, and practical interaction quality"],
            ["Logcat", "Debug Utility", "Tracked runtime logs, errors, and warning messages"],
            ["Firebase Test Lab", "Cloud Testing", "Allowed broader device coverage for release validation"],
            ["Crashlytics", "Post-release QA", "Reported production crashes with diagnostic context"],
            ["Performance Monitoring", "Runtime QA", "Measured latency and performance bottlenecks"],
        ],
        column_widths=[1.75, 1.55, 3.0],
    )

    add_note_box(
        doc,
        "Quality Focus Areas for Lumiere",
        [
            "Screen consistency across splash, onboarding, login, home, bookings, maps, gifting, and profile",
            "Booking and gifting flow correctness",
            "Stable Firebase communication and safe error handling",
            "Performance quality during image loading and navigation",
        ],
    )


def add_inventory_page(doc):
    add_heading(doc, "9. Consolidated Technology Inventory and Conclusion")
    add_body(doc, "The Lumiere portfolio can be summarized as a layered technology stack in which design tools shaped the interface, Android technologies built the mobile client, Firebase powered the backend, AI services accelerated execution, and release tooling supported deployment and maintenance.")
    add_body(doc, "Together, these components formed a coherent delivery pipeline for a salon booking and gifting application that could move from concept to production without relying on unmanaged infrastructure.")

    add_table(
        doc,
        ["Layer", "Primary Components Used", "Status in Assumed Final System"],
        [
            ["Design", "Figma, Figma AI", "Completed and handed off to development"],
            ["Frontend", "Android Studio, Kotlin, XML, Material Components, Glide", "Implemented in the Android app"],
            ["Backend", "Firebase Authentication, Firestore, Cloud Functions, Storage, FCM", "Operational in production"],
            ["Security", "Rules, App Check, validation logic", "Enforced for protected access"],
            ["Deployment", "App Distribution, Play Console, GitHub, Crashlytics", "Used for release and maintenance"],
            ["AI Assistance", "Gemini, ChatGPT/Codex, AI writing support", "Used throughout planning and execution"],
        ],
        column_widths=[1.1, 3.3, 1.9],
    )

    add_note_box(
        doc,
        "Final Conclusion",
        [
            "Lumiere was not defined by a single tool; it was delivered through a connected ecosystem.",
            "Firebase provided the assumed production backend foundation requested for this report.",
            "AI services and modern development tools improved speed, clarity, and consistency across the project lifecycle.",
        ],
    )

    add_body(doc, "Therefore, the completed Lumiere system can be presented as a modern mobile application project built through the coordinated use of Android development utilities, Firebase cloud services, design platforms, deployment tools, testing systems, and AI-enabled assistants.")


def main():
    doc = Document()
    configure_document(doc)

    make_cover(doc)
    add_page_break(doc)
    add_project_overview(doc)
    add_page_break(doc)
    add_frontend_page(doc)
    add_page_break(doc)
    add_firebase_backend_page(doc)
    add_page_break(doc)
    add_data_and_security_page(doc)
    add_page_break(doc)
    add_deployment_page(doc)
    add_page_break(doc)
    add_design_tools_page(doc)
    add_page_break(doc)
    add_ai_tools_page(doc)
    add_page_break(doc)
    add_qa_page(doc)
    add_page_break(doc)
    add_inventory_page(doc)

    for section in doc.sections:
        add_footer(section)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
